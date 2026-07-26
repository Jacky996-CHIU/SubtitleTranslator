"""
Output writers: SRT subtitles, burned-in (hardcoded) video, and Word document.
"""

from __future__ import annotations

import os
import shutil
import subprocess
from typing import List, Optional

from .extractor import Segment
from .runtime_paths import find_tool, imageio_ffmpeg


# --------------------------------------------------------------------------- #
# Timestamps
# --------------------------------------------------------------------------- #
def _fmt_ts(seconds: float, sep: str = ",") -> str:
    if seconds < 0:
        seconds = 0
    ms = int(round(seconds * 1000))
    h, ms = divmod(ms, 3600_000)
    m, ms = divmod(ms, 60_000)
    s, ms = divmod(ms, 1000)
    return f"{h:02d}:{m:02d}:{s:02d}{sep}{ms:03d}"


# --------------------------------------------------------------------------- #
# SRT
# --------------------------------------------------------------------------- #
def write_srt(
    segments: List[Segment],
    path: str,
    mode: str = "translation",   # 'original' | 'translation' | 'bilingual'
) -> str:
    lines = []
    for s in segments:
        if mode == "original":
            body = s.text
        elif mode == "bilingual":
            body = (s.translation + "\n" + s.text).strip()
        else:
            body = s.translation or s.text
        lines.append(str(s.index))
        lines.append(f"{_fmt_ts(s.start)} --> {_fmt_ts(s.end)}")
        lines.append(body.replace("\n", "\n"))
        lines.append("")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return path


# --------------------------------------------------------------------------- #
# ASS (for nicely styled burn-in that mimics the original bold-white caption)
# --------------------------------------------------------------------------- #
def _ass_ts(seconds: float) -> str:
    return _fmt_ts(seconds, sep=".")[:-1]  # ASS uses H:MM:SS.cc (centiseconds)


def write_ass(
    segments: List[Segment],
    path: str,
    mode: str = "translation",
    font: str = "Arial",
    font_size: int = 22,
    play_res_x: int = 640,
    play_res_y: int = 360,
    primary: str = "&H00FFFFFF",   # white (AABBGGRR)
    outline: str = "&H00000000",   # black
    bold: int = -1,                # -1 = true
    alignment: int = 1,            # 1=bottom-left (matches source captions)
    margin_l: int = 24,
    margin_v: int = 28,
    border_style: int = 1,         # 1=outline+shadow, 3=opaque box (masks original)
    box_color: str = "&H00101010", # near-black box (used when border_style=3)
    outline_w: int = 2,
    position_at_original: bool = True,
    shadow: int = 1,
    force_bottom_top: bool = False,
) -> str:
    """Write an ASS subtitle file.

    With ``position_at_original`` each caption is drawn centred on the box where
    the original burned-in caption was detected (``Segment.orig_box``), and its
    font is scaled to that box, so the translation replaces the original in
    place instead of always sitting at the bottom of the frame. Segments without
    a detected box fall back to the global style position.
    """
    header = f"""[Script Info]
ScriptType: v4.00+
PlayResX: {play_res_x}
PlayResY: {play_res_y}
WrapStyle: 2
ScaledBorderAndShadow: yes

[V4+ Styles]
Format: Name, Fontname, Fontsize, PrimaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding
Style: Caption,{font},{font_size},{primary},{outline},{box_color},{bold},0,0,0,100,100,0,0,{border_style},{outline_w},{shadow},{alignment},{margin_l},20,{margin_v},1

[Events]
Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text
"""
    ev = []
    for s in segments:
        if mode == "original":
            body = s.text
        elif mode == "bilingual":
            body = (s.translation + "\n" + s.text).strip()
        else:
            body = s.translation or s.text
        body = body.replace("\n", "\\N")

        tags = ""
        if position_at_original and s.orig_box and not force_bottom_top:
            bx, by, bw, bh = s.orig_box
            cx, cy = int(bx + bw / 2), int(by + bh / 2)
            nlines = body.count("\\N") + 1
            # Scale the text to the original caption's height so it reads like a
            # replacement rather than an overlay.
            fs = max(12, min(int(bh / max(1, nlines) * 0.78), play_res_y))
            tags = f"{{\\an5\\pos({cx},{cy})\\fs{fs}}}"

        ev.append(
            f"Dialogue: 0,{_ass_ts(s.start)},{_ass_ts(s.end)},Caption,,0,0,0,,{tags}{body}"
        )
    with open(path, "w", encoding="utf-8") as f:
        f.write(header + "\n".join(ev) + "\n")
    return path


# --------------------------------------------------------------------------- #
# Burn subtitles into video (ffmpeg)
# --------------------------------------------------------------------------- #
def _ffmpeg_candidates() -> List[str]:
    """All ffmpeg binaries we could use, best-known-first."""
    cands = [
        os.environ.get("FFMPEG_BINARY"),
        imageio_ffmpeg(),        # pip wheel: static build WITH libass
        find_tool("ffmpeg"),     # bundled next to the frozen app
        shutil.which("ffmpeg"),  # system install
    ]
    seen, out = set(), []
    for c in cands:
        if c and c not in seen:
            seen.add(c)
            out.append(c)
    return out


def ffmpeg_path() -> str:
    cands = _ffmpeg_candidates()
    return cands[0] if cands else "ffmpeg"


def _has_subtitles_filter(ffmpeg: str) -> bool:
    """True if this ffmpeg build can render subtitles (i.e. was built w/ libass).

    Homebrew/distro builds are sometimes compiled without libass, in which case
    the ``subtitles`` filter is missing entirely and burning text is impossible.
    """
    try:
        out = subprocess.run([ffmpeg, "-hide_banner", "-filters"],
                             capture_output=True, text=True, timeout=30).stdout
    except Exception:
        return False
    return any(line.split()[1:2] == ["subtitles"]
               for line in out.splitlines() if line.strip())


_SUBS_FFMPEG_CACHE: dict = {}


def ffmpeg_with_subtitles() -> str:
    """An ffmpeg that can actually burn subtitles, or raise a clear error."""
    if "path" in _SUBS_FFMPEG_CACHE:
        return _SUBS_FFMPEG_CACHE["path"]
    tried = []
    for cand in _ffmpeg_candidates():
        tried.append(cand)
        if _has_subtitles_filter(cand):
            _SUBS_FFMPEG_CACHE["path"] = cand
            return cand
    raise RuntimeError(
        "找不到支持字幕渲染的 ffmpeg（缺少 libass / subtitles 滤镜），无法压制视频。\n"
        "No ffmpeg with subtitle support (libass) was found, so the video cannot "
        "be burned.\n"
        "解决 / Fix:  pip install imageio-ffmpeg   或安装带 libass 的 ffmpeg。\n"
        + ("已尝试 / tried: " + ", ".join(tried) if tried else "")
    )


def ffprobe_path() -> str:
    return (os.environ.get("FFPROBE_BINARY")
            or find_tool("ffprobe")           # bundled binary when frozen
            or shutil.which("ffprobe")
            or ffmpeg_path().replace("ffmpeg", "ffprobe"))


def _probe_dims(video_in: str):
    """Video pixel dimensions. Uses OpenCV first — ffprobe is not shipped by
    every ffmpeg distribution (e.g. the imageio-ffmpeg wheel), and a wrong
    fallback size would misplace the cover boxes."""
    try:
        import cv2
        cap = cv2.VideoCapture(video_in)
        w = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        h = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        cap.release()
        if w > 0 and h > 0:
            return w, h
    except Exception:
        pass
    try:
        out = subprocess.check_output([
            ffprobe_path(),
            "-v", "error", "-select_streams", "v:0",
            "-show_entries", "stream=width,height",
            "-of", "csv=p=0:s=x", video_in,
        ], text=True).strip()
        w, h = out.split("x")[:2]
        return int(w), int(h)
    except Exception:
        return 640, 360


def _cover_filters(segments: List[Segment], w: int, h: int, pad: int = 6) -> str:
    """drawbox filters that mask each original caption for its time range."""
    parts = []
    for s in segments:
        if not s.orig_box:
            continue
        x, y, bw, bh = s.orig_box
        x = max(0, x - pad); y = max(0, y - pad)
        bw = min(w - x, bw + 2 * pad); bh = min(h - y, bh + 2 * pad)
        st, en = s.start, s.end
        parts.append(
            f"drawbox=x={x}:y={y}:w={bw}:h={bh}:color=black@1.0:t=fill:"
            f"enable='between(t,{st:.2f},{en:.2f})'"
        )
    return ",".join(parts)


def burn_video(
    video_in: str,
    segments: List[Segment],
    video_out: str,
    workdir: str,
    mode: str = "translation",
    font: str = "Arial",
    font_size: Optional[int] = None,
    fonts_dir: Optional[str] = None,
    cover_original: bool = True,
    progress_cb=None,
    style=None,
    audio_from: Optional[str] = None,
    encoder: Optional[str] = None,
) -> str:
    """Render a new MP4 with the (translated) captions burned in.

    When ``cover_original`` is True, the original burned-in caption is masked
    with a fill box (using the detected bounding box) before the translated
    caption is drawn, so the result shows only the translation.

    ``audio_from`` supplies the audio track from a different file, which the
    inpainting path needs because its intermediate video carries no audio.
    ``style`` is a :class:`subtrans.style.SubtitleStyle` overriding appearance.
    """
    w, h = _probe_dims(video_in)
    ass_file = os.path.join(workdir, "burn.ass")

    if style is not None:
        fs = style.resolved_font_size(h)
        write_ass(segments, ass_file, mode=mode,
                  font=style.font or font, font_size=fs,
                  play_res_x=w, play_res_y=h,
                  primary=style.primary_ass, outline=style.outline_ass,
                  box_color=style.back_ass, border_style=style.border_style,
                  outline_w=style.outline_width, shadow=style.shadow,
                  bold=-1 if style.bold else 0,
                  alignment=style.alignment, margin_v=style.margin_v,
                  position_at_original=(style.position == "original"),
                  force_bottom_top=(style.position != "original"))
    else:
        fs = font_size or max(18, int(h * 0.062))
        write_ass(segments, ass_file, mode=mode, font=font, font_size=fs,
                  play_res_x=w, play_res_y=h)

    # Run ffmpeg from the workdir and reference the subtitle file by its bare
    # name. A relative name has no "/" or ":" so it sidesteps ffmpeg's brittle
    # filtergraph path-escaping (which broke on absolute paths / hidden dirs).
    ass_name = os.path.basename(ass_file)
    fonts_opt = ""
    if fonts_dir:
        fd = os.path.abspath(fonts_dir).replace("\\", "/").replace(":", "\\:")
        fonts_opt = f":fontsdir='{fd}'"

    # Different ffmpeg builds disagree on how the subtitles filter accepts its
    # filename (some require the explicit ``filename=`` key and reject the bare
    # positional form with "No option name near ..."). Try the known-good forms
    # in order until one succeeds, so it works across ffmpeg versions.
    sub_variants = [
        f"subtitles=filename={ass_name}{fonts_opt}",
        f"subtitles={ass_name}{fonts_opt}",
        f"subtitles='{ass_name}'{fonts_opt}",
    ]

    cover = _cover_filters(segments, w, h) if cover_original else ""

    ffmpeg = ffmpeg_with_subtitles()   # raises a clear error if none can burn

    vcodec = encoder or "libx264"
    last_err = ""
    for sub in sub_variants:
        vf = (cover + "," + sub) if cover else sub
        cmd = [ffmpeg, "-y", "-i", os.path.abspath(video_in)]
        if audio_from:
            # Video from the (inpainted) input, audio from the original file.
            cmd += ["-i", os.path.abspath(audio_from),
                    "-map", "0:v:0", "-map", "1:a:0?", "-shortest"]
        cmd += ["-vf", vf, "-c:v", vcodec]
        if vcodec == "libx264":
            cmd += ["-preset", "medium", "-crf", "20"]
        else:
            cmd += ["-b:v", "0", "-q:v", "50"]     # hardware encoders
        cmd += ["-c:a", "copy", os.path.abspath(video_out)]
        proc = subprocess.run(cmd, capture_output=True, text=True, cwd=workdir)
        if proc.returncode == 0:
            return video_out
        last_err = proc.stderr

    raise RuntimeError("ffmpeg burn failed:\n" + last_err[-2000:])


# --------------------------------------------------------------------------- #
# Word document
# --------------------------------------------------------------------------- #
def write_docx(
    segments: List[Segment],
    path: str,
    title: str = "字幕翻译 / Subtitle Translation",
    source_lang: str = "EN",
    target_lang: str = "",
    video_name: str = "",
) -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(title, level=0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"视频 / Video: {video_name}\n"
        f"源语言 / Source: {source_lang}    目标语言 / Target: {target_lang}\n"
        f"字幕条数 / Segments: {len(segments)}"
    ).font.size = Pt(10)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for c, t in zip(hdr, ["#", "时间 / Time", "原文 / Original", "译文 / Translation"]):
        run = c.paragraphs[0].add_run(t)
        run.bold = True

    for s in segments:
        row = table.add_row().cells
        row[0].text = str(s.index)
        row[1].text = f"{_fmt_ts(s.start)}\n{_fmt_ts(s.end)}"
        row[2].text = s.text.replace("\n", " ")
        row[3].text = (s.translation or "").replace("\n", " ")

    doc.save(path)
    return path
